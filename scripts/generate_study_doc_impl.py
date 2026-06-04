from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = REPO_ROOT / "docs" / "WARELYN_PROJECT_STUDY_GUIDE.docx"
TODAY = date.today().strftime("%d %B %Y")

PRIMARY = "1F3A5F"
ACCENT = "1A7F5A"
TABLE_HEADER = "2E4F7A"
ALT_ROW = "F0F4F8"
LIGHT_BLUE = "D6E4F0"
LIGHT_GREEN = "D4EDDA"
LIGHT_AMBER = "FFF3CD"
LIGHT_RED = "F8D7DA"
BORDER = "CCCCCC"
TEXT = "333333"
MUTED = "666666"
WHITE = "FFFFFF"


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.replace("#", "")
    return RGBColor.from_string(value)


def set_run_font(run, *, name="Calibri", size=11, bold=False, italic=False, color=TEXT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)


def set_paragraph_format(paragraph, *, before=0, after=6, line=1.15, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def set_style_font(style, *, name, size, bold=False, color=TEXT, italic=False):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = rgb(color)
    if style._element.rPr is None:
        style._element.get_or_add_rPr()
    rpr = style._element.rPr
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)


def add_paragraph(doc, text="", *, style="Normal", align=None, bold=False, italic=False, size=None, color=TEXT, before=0, after=6, line=1.15):
    p = doc.add_paragraph(style=style)
    set_paragraph_format(p, before=before, after=after, line=line, align=align)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size or 11, bold=bold, italic=italic, color=color)
    return p


def add_mixed_paragraph(doc, parts, *, style="Normal", align=None, before=0, after=6, line=1.15):
    p = doc.add_paragraph(style=style)
    set_paragraph_format(p, before=before, after=after, line=line, align=align)
    for text, opts in parts:
        run = p.add_run(text)
        set_run_font(
            run,
            name=opts.get("name", "Calibri"),
            size=opts.get("size", 11),
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
            color=opts.get("color", TEXT),
        )
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    if level == 1:
        set_paragraph_format(p, before=14, after=6, line=1.05)
    elif level == 2:
        set_paragraph_format(p, before=10, after=4, line=1.05)
    else:
        set_paragraph_format(p, before=8, after=3, line=1.05)
    return p


def add_bullets(doc, items, *, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if level:
            p.paragraph_format.left_indent = Cm(0.5 * level)
        set_paragraph_format(p, before=0, after=2, line=1.1)
        run = p.add_run(item)
        set_run_font(run)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_format(p, before=0, after=2, line=1.1)
        run = p.add_run(item)
        set_run_font(run)


def set_cell_text(cell, text="", *, bold=False, size=10.5, color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_format(p, before=0, after=0, line=1.05)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tcMar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tcMar.append(el)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_table_widths(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col_idx, width in enumerate(widths_cm):
        for cell in table.columns[col_idx].cells:
            cell.width = Cm(width)


def build_table(doc, headers, rows, widths_cm, *, alignments=None, font_size=10.0, header_fill=TABLE_HEADER, alt_fill=ALT_ROW):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    set_table_widths(table, widths_cm)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, size=font_size, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        set_cell_border(
            cell,
            top={"val": "single", "sz": 8, "space": 0, "color": BORDER},
            bottom={"val": "single", "sz": 8, "space": 0, "color": BORDER},
            left={"val": "single", "sz": 8, "space": 0, "color": BORDER},
            right={"val": "single", "sz": 8, "space": 0, "color": BORDER},
        )
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT
            if alignments and c_idx < len(alignments):
                align = alignments[c_idx]
            set_cell_text(cells[c_idx], str(value), size=font_size, align=align)
            if r_idx % 2 == 1:
                set_cell_shading(cells[c_idx], alt_fill)
            set_cell_margins(cells[c_idx])
            set_cell_border(
                cells[c_idx],
                top={"val": "single", "sz": 4, "space": 0, "color": BORDER},
                bottom={"val": "single", "sz": 4, "space": 0, "color": BORDER},
                left={"val": "single", "sz": 4, "space": 0, "color": BORDER},
                right={"val": "single", "sz": 4, "space": 0, "color": BORDER},
            )
    return table


def build_callout(doc, lines, fill=LIGHT_BLUE, title=None):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=140, bottom=120, end=140)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 8, "space": 0, "color": BORDER},
        bottom={"val": "single", "sz": 8, "space": 0, "color": BORDER},
        left={"val": "single", "sz": 8, "space": 0, "color": BORDER},
        right={"val": "single", "sz": 8, "space": 0, "color": BORDER},
    )
    cell.text = ""
    if title:
        add_mixed_paragraph(
            cell.add_paragraph(),
            [(title, {"bold": True, "color": PRIMARY, "size": 11})],
            before=0,
            after=3,
            line=1.05,
        )
    for line in lines:
        p = cell.add_paragraph()
        set_paragraph_format(p, before=0, after=2, line=1.1)
        run = p.add_run(line)
        set_run_font(run, size=10.5)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p


def add_field(paragraph, field_code: str, placeholder: str = "1"):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)


def add_toc(paragraph):
    add_field(paragraph, r'TOC \o "1-2" \h \z \u', "Table of contents will appear here when fields are updated.")


def add_page_number(paragraph):
    add_field(paragraph, "PAGE", "1")


def apply_document_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, name="Calibri", size=11, color=TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for level, size, color, bold in [("Heading 1", 18, PRIMARY, True), ("Heading 2", 14, ACCENT, True), ("Heading 3", 12, "2E4F7A", True)]:
        style = doc.styles[level]
        set_style_font(style, name="Arial", size=size, bold=bold, color=color)
    doc.styles["List Bullet"].font.name = "Calibri"
    doc.styles["List Bullet"].font.size = Pt(11)
    doc.styles["List Number"].font.name = "Calibri"
    doc.styles["List Number"].font.size = Pt(11)


def configure_sections(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    section.different_first_page_header_footer = True
    return section


def table_without_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
            )


def add_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    header_table = header.add_table(rows=1, cols=2, width=Cm(16))
    table_without_borders(header_table)
    header_table.columns[0].width = Cm(12)
    header_table.columns[1].width = Cm(4)
    set_cell_text(header_table.cell(0, 0), "Warelyn Inventory — Project Study Guide", size=9, color=MUTED)
    set_cell_text(header_table.cell(0, 1), "CONFIDENTIAL", size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for cell in header_table.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_table = footer.add_table(rows=1, cols=3, width=Cm(16))
    table_without_borders(footer_table)
    footer_table.columns[0].width = Cm(7.5)
    footer_table.columns[1].width = Cm(1.0)
    footer_table.columns[2].width = Cm(7.5)
    set_cell_text(footer_table.cell(0, 0), "© 2025 Warelyn Inventory", size=8.5, color=MUTED)
    center = footer_table.cell(0, 1)
    center.text = ""
    p = center.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=0, line=1.0)
    add_page_number(p)
    set_run_font(p.runs[0], size=8.5, color=MUTED)
    set_cell_text(footer_table.cell(0, 2), TODAY, size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for cell in footer_table.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)


def add_cover_page(doc):
    logo_table = doc.add_table(rows=1, cols=1)
    logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo_cell = logo_table.cell(0, 0)
    logo_cell.width = Cm(8.5)
    set_cell_shading(logo_cell, LIGHT_BLUE)
    set_cell_border(
        logo_cell,
        top={"val": "single", "sz": 10, "space": 0, "color": PRIMARY},
        bottom={"val": "single", "sz": 10, "space": 0, "color": PRIMARY},
        left={"val": "single", "sz": 10, "space": 0, "color": PRIMARY},
        right={"val": "single", "sz": 10, "space": 0, "color": PRIMARY},
    )
    set_cell_margins(logo_cell, top=200, start=160, bottom=200, end=160)
    logo_cell.text = ""
    p = logo_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("WARELYN")
    set_run_font(run, name="Arial", size=84, bold=True, color=PRIMARY)
    set_paragraph_format(p, before=0, after=0, line=1.0)

    add_paragraph(doc, "", after=10)
    add_paragraph(doc, "Warelyn Inventory", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=28, color=PRIMARY, before=10, after=0, line=1.0)
    add_paragraph(
        doc,
        "Multi-Tenant Inventory SaaS — Complete Project Study Guide",
        style="Normal",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=15,
        color=ACCENT,
        before=0,
        after=2,
        line=1.0,
    )
    add_paragraph(doc, "Version 1.1 · June 2025", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=MUTED, after=8)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(rule, before=0, after=16, line=1.0)
    run = rule.add_run(" ")
    set_run_font(run, size=1)
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "24")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), PRIMARY)
    pBdr.append(bottom)
    pPr.append(pBdr)

    intro = (
        "This document is a comprehensive end-to-end reference for the Warelyn Inventory platform. "
        "It covers the system architecture, database design, API contracts, user roles and permissions, "
        "business workflows, AI features, deployment, and development standards. It is intended for developers "
        "joining the project, technical stakeholders, and anyone who needs a complete understanding of how Warelyn works."
    )
    add_paragraph(doc, intro, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=10, line=1.18)
    add_page_break(doc)


def add_toc_page(doc):
    add_heading(doc, "Table of Contents", 1)
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=6, line=1.05)
    add_toc(p)
    add_paragraph(doc, "Update this table of contents by pressing Ctrl+A then F9 in Microsoft Word.", color=MUTED, italic=True, after=12)
    add_page_break(doc)


def section_overview(doc):
    add_heading(doc, "3. Project Overview", 1)
    add_heading(doc, "3.1 What is Warelyn?", 2)
    for para in [
        "Warelyn Inventory is a multi-tenant inventory management SaaS for wholesale and retail businesses that need accurate stock visibility, structured workflows, and reliable operational control.",
        "Its core value proposition is role-aware workflow automation: when one role completes a step, the next step is automatically routed to the correct person.",
        "The central product question throughout the codebase is simple but critical: when a role completes a workflow step, does the next step automatically move to the right person?",
        "Each tenant operates in isolated data space with its own users, settings, currency, reports, and documents.",
    ]:
        add_paragraph(doc, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "3.2 Technology Stack", 2)
    tech_rows = [
        ("Backend Language", "Python 3.12+"),
        ("Backend Framework", "FastAPI 0.115.6"),
        ("ORM", "SQLAlchemy 2.0.45"),
        ("Database Migrations", "Alembic 1.13.3"),
        ("Data Validation", "Pydantic V2 (pydantic-settings 2.7.1)"),
        ("Authentication", "PyJWT 2.10.1 + bcrypt 5.0.0"),
        ("PDF Generation", "WeasyPrint 68.1"),
        ("Template Engine", "Jinja2 3.1.4+"),
        ("Email", "SMTP (configurable) + MailHog for dev"),
        ("AI / LLM", "Google Gemini 2.5 Flash"),
        ("AI Embeddings", "Gemini embedding model"),
        ("Vector / Document DB", "MongoDB (pymongo 4.10.1)"),
        ("Relational DB", "MySQL (PyMySQL 1.1.1)"),
        ("Rate Limiting", "SlowAPI 0.1.9"),
        ("Testing", "pytest 8.3.4"),
        ("Frontend Framework", "React + Vite"),
        ("UI Styling", "Tailwind CSS 3.4.17"),
        ("Charts", "Recharts 2.15.3"),
        ("Icons", "Lucide React 1.16.0"),
        ("Containerisation", "Docker + Docker Compose"),
    ]
    build_table(doc, ["Component", "Technology + Version"], tech_rows, [5.0, 11.0], font_size=9.2, alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    add_heading(doc, "3.3 Key Features Summary", 2)
    add_bullets(doc, [
        "Multi-tenant architecture with complete data isolation.",
        "Role-based access control for SUPER_ADMIN, TENANT_ADMIN, INVENTORY_MANAGER, SALES_STAFF, PURCHASE_STAFF, and VIEWER.",
        "Full sales order lifecycle from draft to invoice.",
        "Full purchase order lifecycle from draft to bill.",
        "Customer returns with QC inspection workflow.",
        "Inventory engine with stock ledger, reservations, batch tracking, serial tracking, and expiry tracking.",
        "Cycle counting and stock reconciliation.",
        "Workflow task engine with automatic task creation and role handoff.",
        "In-app notification pipeline with role-targeted delivery.",
        "AI Copilot for natural language queries over live tenant data.",
        "FAQ Assistant powered by RAG and hybrid search.",
        "Role-specific dashboards with real-time KPIs and charts.",
        "Document templates for invoice and bill PDF/email generation.",
        "Multi-currency support across 30+ ISO 4217 currencies.",
        "CSV product import.",
        "Audit logs for significant actions.",
        "Forgot password via OTP email.",
    ])
    add_page_break(doc)


def section_architecture(doc):
    add_heading(doc, "4. System Architecture", 1)
    add_heading(doc, "4.1 High-Level Architecture", 2)
    for para in [
        "The backend is a FastAPI REST API with a strict 5-layer architecture. Models hold persistence shape, repositories handle DB access, services own business logic and commits, API routers translate HTTP, and schemas define request and response contracts.",
        "The frontend is a React single-page application that communicates only through the REST API. Tenant data isolation is enforced at the database-query level, and Docker Compose orchestrates the services for local development.",
    ]:
        add_paragraph(doc, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    build_table(doc, ["Layer", "Directory", "Responsibility", "Rules"], [
        ("Models", "backend/app/models/", "SQLAlchemy ORM definitions, persistence shape", "No business logic; no imports from services"),
        ("Repositories", "backend/app/repositories/", "Database access only, always filtered by tenant_id", "No db.commit(); no business logic"),
        ("Services", "backend/app/services/", "Business logic, orchestration, db.commit() ownership", "Calls repositories and external services; never called by repositories"),
        ("API (Routers)", "backend/app/api/", "HTTP translation only, calls services", "require_roles() on every endpoint; returns schema types"),
        ("Schemas", "backend/app/schemas/", "Pydantic request/response contracts", "Never contains ORM objects; never used in repositories"),
    ], [2.2, 4.5, 5.2, 4.1], font_size=8.8)

    add_heading(doc, "4.2 Request Lifecycle", 2)
    add_numbers(doc, [
        "HTTP request arrives with a Bearer token in the Authorization header.",
        "HTTPBearer extracts the token.",
        "get_current_user_context() decodes the JWT and loads UserContext (user, role, tenant_id).",
        "require_roles(*allowed_roles) checks the role and returns 403 if not allowed.",
        "The API route calls the appropriate Service method, passing tenant_id.",
        "The Service validates business rules and calls Repository methods.",
        "The Repository executes SQL queries, always filtered by tenant_id.",
        "The Service calls db.commit() to persist the transaction.",
        "Side effects fire in try/except: workflow task creation, notification, and audit log.",
        "The Service returns the ORM object, and the API route serialises it to a schema.",
        "The HTTP response returns with an X-Request-ID header.",
    ])

    add_heading(doc, "4.3 Frontend Architecture", 2)
    build_table(doc, ["Layer", "Directory", "Responsibility"], [
        ("Pages", "pages/", "Thin screens that compose UI, call services and hooks. One useEffect per page load."),
        ("Services", "services/", "All API calls. One file per domain. Uses shared apiRequest() helper."),
        ("Components", "components/", "Reusable UI components with zero business logic."),
        ("UI Primitives", "components/ui/", "Button, Input, Badge, Card, Table, Modal, and similar atoms."),
        ("Context", "context/", "Global state such as AuthContext and TenantSettingsContext."),
        ("Hooks", "hooks/", "Shared stateful logic used by 2+ pages."),
        ("Routes", "routes/AppRoutes.jsx", "All routes with RoleGuard wrappers and lazy loading."),
        ("Navigation", "components/navigation.js", "Sidebar items with role visibility and detailRoutes for breadcrumbs."),
    ], [2.1, 4.0, 9.9], font_size=8.8)

    add_heading(doc, "4.4 Multi-Tenancy Model", 2)
    for para in [
        "Tenant isolation works by ensuring every tenant-owned table carries a tenant_id column and every repository query filters by that tenant_id. This keeps one tenant's stock, orders, reports, and documents separated from another tenant's data.",
        "The tenant_id value is sourced from the authenticated JWT payload after login and is validated on every request through UserContext. It is never accepted from the request body for access control.",
        "The TenantScopedRepository pattern centralises the filtering rule so repositories cannot accidentally forget the tenant boundary. That makes isolation the default rather than an optional check that developers must remember to add each time.",
    ]:
        add_paragraph(doc, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    build_callout(doc, [
        "Security note: The backend is the security boundary. Frontend role hiding is UX only — every API endpoint enforces its own require_roles() check independently.",
    ], fill=LIGHT_BLUE)
    add_page_break(doc)


def section_database(doc):
    add_heading(doc, "5. Database Design", 1)
    add_heading(doc, "5.1 Database Overview", 2)
    add_paragraph(doc, "Warelyn uses a MySQL relational database. The project uses many Alembic migrations to model the full entity system. Every tenant-owned table has a tenant_id foreign key to the tenants table, and all status columns are stored as VARCHAR values with application-level enums for migration safety.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "5.2 Core Entity Tables", 2)
    entity_groups = [
        ("Auth & Tenants", [
            ("tenants", "id, company_name, contact_email, currency, timezone, status", "One row per business customer"),
            ("users", "id, tenant_id, name, email, role, status, password_hash", "Users within a tenant; SUPER_ADMIN has no tenant_id"),
            ("refresh_tokens", "id, user_id, token_hash, expires_at, revoked", "Refresh token store for JWT rotation"),
        ]),
        ("Catalog (Master Data)", [
            ("categories", "id, tenant_id, name", "Product categories"),
            ("brands", "id, tenant_id, name", "Product brands"),
            ("vendors", "id, tenant_id, name, contact_email", "Suppliers"),
            ("customers", "id, tenant_id, name, contact_email", "Buyers"),
            ("products", "id, tenant_id, sku, name, category_id, brand_id, tracking_type, cost_price, selling_price, reorder_level", "Core product catalog"),
        ]),
        ("Inventory Engine", [
            ("warehouses", "id, tenant_id, name, address", "Physical warehouses"),
            ("warehouse_locations", "id, tenant_id, warehouse_id, name", "Locations or bins inside warehouses"),
            ("warehouse_stock", "id, tenant_id, product_id, warehouse_id, quantity_on_hand, quantity_reserved", "Current stock projection"),
            ("stock_ledger_entries", "id, tenant_id, product_id, warehouse_id, location_id, quantity_delta, movement_type, reference_type, reference_id", "Immutable stock history"),
            ("stock_reservations", "id, tenant_id, product_id, warehouse_id, quantity_reserved, reference_type, reference_id, status", "Active reservations against orders"),
            ("inventory_batches", "id, tenant_id, product_id, warehouse_id, batch_number, expiry_date, quantity", "Batch-tracked stock groups"),
            ("inventory_serials", "id, tenant_id, product_id, serial_number, status", "Individual serial numbers"),
            ("blocked_return_stock", "id, tenant_id, product_id, warehouse_id, quantity, block_reason", "Stock blocked pending disposition"),
        ]),
        ("Sales", [
            ("sales_orders", "id, tenant_id, order_number, customer_id, status, created_by", "Sales order header"),
            ("sales_order_items", "id, tenant_id, sales_order_id, product_id, ordered_quantity, fulfilled_quantity, unit_price", "Sales order line items"),
            ("sales_fulfillments", "id, tenant_id, sales_order_id, status", "Fulfillment header"),
            ("packages", "id, tenant_id, sales_order_id, package_number, status", "Packed package"),
            ("pick_tasks", "id, tenant_id, sales_order_id, status, created_by", "Pick task per order"),
        ]),
        ("Purchasing", [
            ("purchase_orders", "id, tenant_id, po_number, vendor_id, status", "Purchase order header"),
            ("purchase_order_items", "id, tenant_id, purchase_order_id, product_id, ordered_quantity, unit_price", "Purchase order line items"),
            ("purchase_receipts", "id, tenant_id, purchase_order_id, status", "Receipt header"),
            ("purchase_receipt_items", "id, tenant_id, purchase_receipt_id, product_id, received_quantity", "Receipt line items"),
        ]),
        ("Returns", [
            ("sales_returns", "id, tenant_id, sales_order_id, return_number, status", "Return header"),
            ("sales_return_items", "id, tenant_id, sales_return_id, product_id, returned_quantity, qc_status, accepted_quantity, rejected_quantity", "Return line items with QC outcomes"),
            ("return_qc_inspections", "id, tenant_id, sales_return_id, inspector_id, created_at", "Inspection audit record"),
        ]),
        ("Operations", [
            ("putaway_tasks", "id, tenant_id, receipt_id, warehouse_id, status", "Putaway task per receipt"),
            ("stock_count_sessions", "id, tenant_id, warehouse_id, session_number, status", "Cycle count session"),
            ("stock_count_lines", "id, tenant_id, session_id, product_id, location_id, system_quantity, counted_quantity, variance", "Count line per product"),
        ]),
        ("Workflow", [
            ("workflow_events", "id, tenant_id, event_type, entity_type, entity_id, actor_user_id, payload_json", "Domain event log"),
            ("workflow_tasks", "id, tenant_id, step_key, entity_type, entity_id, assigned_role, status, priority, action_url", "Role task queue"),
        ]),
        ("Communication & Documents", [
            ("notifications", "id, tenant_id, user_id, title, message, type, category, is_read, cleared_at", "In-app notifications"),
            ("invoices", "id, tenant_id, invoice_number, sales_order_id, customer_id, currency_code, status", "Customer invoices"),
            ("bills", "id, tenant_id, bill_number, purchase_order_id, vendor_id, currency_code, status", "Vendor bills"),
            ("document_templates", "id, tenant_id, template_key, format, subject, body", "Jinja2 email and PDF templates"),
        ]),
        ("AI Assistant", [
            ("faq_chunks", "id, tenant_id, content, searchable_text, embedding_vector, source_uri", "RAG knowledge chunks"),
            ("assistant_sessions", "id, tenant_id, user_id, title", "Copilot session"),
            ("assistant_messages", "id, tenant_id, session_id, role, content, confidence_score, citations_json", "Copilot messages"),
        ]),
        ("Audit", [
            ("audit_logs", "id, tenant_id, actor_user_id, action, entity_type, entity_id, before_json, after_json", "Full audit trail"),
        ]),
    ]
    for title, rows in entity_groups:
        add_heading(doc, title, 3)
        build_table(doc, ["Table", "Key Columns", "Purpose"], rows, [2.6, 7.5, 5.0], font_size=8.5)

    add_heading(doc, "5.3 Inventory Engine", 2)
    for para in [
        "The InventoryEngine is the only allowed path for stock mutations. No service or router may bypass it, because the engine keeps the stock ledger and the stock projection in sync.",
        "The inventory model is ledger-first: every stock change creates an immutable StockLedgerEntry before the warehouse_stock projection is updated.",
        "Idempotency is mandatory for stock mutations. A duplicate call with the same idempotency key is silently ignored instead of double-applying the movement.",
        "Movement types include STOCK_IN, STOCK_OUT, TRANSFER_IN, TRANSFER_OUT, STOCK_ADJUSTMENT, RETURN_RESTOCK, RETURN_BLOCK, and BATCH_EXPIRED.",
    ]:
        add_paragraph(doc, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    build_table(doc, ["Tracking Type", "Meaning", "Typical Use"], [
        ("STANDARD", "Total quantity tracked", "Clothing, stationery, commodities"),
        ("BATCH", "Grouped by batch number and expiry date", "Food, pharmaceuticals, cosmetics"),
        ("SERIAL", "Individual unit serial numbers", "Electronics, machinery, high-value items"),
    ], [3.0, 6.2, 6.0], font_size=9.2)
    add_page_break(doc)


def role_block(doc, role, who, can_do, cannot_do, tasks):
    add_heading(doc, role, 3)
    add_paragraph(doc, f"Who: {who}", bold=True, before=0, after=2)
    add_bullets(doc, [f"Can: {item}" for item in can_do])
    add_bullets(doc, [f"Cannot: {item}" for item in cannot_do])
    add_paragraph(doc, f"Workflow tasks: {tasks}", bold=True, before=0, after=4)


def section_roles(doc):
    add_heading(doc, "6. User Roles and Permissions", 1)
    add_heading(doc, "6.1 Role Overview", 2)
    add_paragraph(doc, "Warelyn has 6 user roles. The backend enforces all permissions via require_roles() on every API endpoint. The frontend hides actions based on role, but the backend is the true security boundary.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "6.2 Role Definitions", 2)
    role_block(
        doc,
        "SUPER_ADMIN",
        "Warelyn platform administrator — not a tenant user.",
        [
            "manage all tenants",
            "view platform-wide audit logs",
            "view platform health",
            "enable or disable tenants",
            "view the admin dashboard",
        ],
        [
            "access any tenant's business data such as orders, stock, or reports",
            "appear outside /admin/* routes",
        ],
        "No workflow tasks — operates at platform level",
    )
    role_block(
        doc,
        "TENANT_ADMIN",
        "Business owner or operations manager.",
        [
            "manage users and roles",
            "configure company settings and currency",
            "manage document templates",
            "approve high-value purchase orders",
            "view all reports and dashboards",
            "complete any workflow task",
            "view the full audit log",
            "use the AI Copilot",
        ],
        ["access another tenant's data"],
        "APPROVE_PO (high-value purchase orders)",
    )
    role_block(
        doc,
        "INVENTORY_MANAGER",
        "Warehouse manager or stock controller.",
        [
            "manage products, warehouses, and locations",
            "perform stock adjustments, transfers, cycle counts, and reconciliation",
            "manage batch, serial, and expiry data",
            "complete pick tasks and putaway tasks",
            "inspect and process returns",
            "view inventory and warehouse reports",
            "manage document templates",
        ],
        [
            "manage users",
            "change company settings",
            "create or manage sales orders or purchase orders",
        ],
        "PICK_ORDER, PUTAWAY_STOCK, RETURN_QC",
    )
    role_block(
        doc,
        "SALES_STAFF",
        "Sales representative or customer success user.",
        [
            "create and confirm sales orders",
            "cancel sales orders",
            "manage customers",
            "create and send invoices",
            "create and submit sales returns",
            "view fulfillment status",
        ],
        [
            "inspect or process returns",
            "create or manage purchase orders",
            "manage products or warehouses",
            "manage users",
        ],
        "CREATE_INVOICE",
    )
    role_block(
        doc,
        "PURCHASE_STAFF",
        "Procurement officer or buyer.",
        [
            "create and submit purchase orders",
            "commit purchase receipts",
            "record vendor bills",
            "manage vendors",
            "view purchase and supplier reports",
        ],
        [
            "create or manage sales orders",
            "create invoices",
            "manage customers",
            "inspect returns",
            "manage users",
        ],
        "RECORD_BILL, REORDER_STOCK, APPROVE_PO (when directly assigned)",
    )
    role_block(
        doc,
        "VIEWER",
        "Read-only stakeholder such as an executive or auditor.",
        [
            "view all reports",
            "view the dashboard in read-only mode",
            "view the product catalog",
            "view stock levels",
        ],
        [
            "create, edit, or delete anything",
            "perform any workflow action",
        ],
        "No workflow tasks",
    )

    add_heading(doc, "6.3 RBAC Enforcement Model", 2)
    build_table(doc, ["Endpoint", "TENANT_ADMIN", "INVENTORY_MANAGER", "SALES_STAFF", "PURCHASE_STAFF", "VIEWER"], [
        ("POST /api/sales-orders", "✅", "✅", "✅", "❌", "❌"),
        ("POST /api/purchases", "✅", "✅", "❌", "✅", "❌"),
        ("POST /api/returns", "✅", "✅", "✅", "❌", "❌"),
        ("POST /api/returns/{id}/inspect", "✅", "✅", "❌", "❌", "❌"),
        ("GET /api/reports/*", "✅", "✅", "❌", "❌", "✅"),
        ("POST /api/workflow/tasks/{id}/complete", "✅", "role match", "role match", "role match", "❌"),
        ("GET /api/settings/users", "✅", "❌", "❌", "❌", "❌"),
        ("GET /admin/*", "SUPER_ADMIN only", "❌", "❌", "❌", "❌"),
    ], [4.0, 2.2, 2.4, 2.2, 2.3, 1.8], font_size=8.4, alignments=[WD_ALIGN_PARAGRAPH.LEFT]*6)
    add_page_break(doc)


def section_workflows(doc):
    add_heading(doc, "7. Business Workflows", 1)
    add_paragraph(doc, "Every major business action in Warelyn follows the same chain: action → permission check → entity status change → domain event logged → workflow task created → notification sent → audit log written. Steps after the status change are side effects — they run in try/except and never break the primary operation.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "7.1 Sales Order Workflow", 2)
    build_table(doc, ["Step", "Action", "Actor", "Status Change", "Task Created", "Notification Sent"], [
        ("1", "Create sales order", "SALES_STAFF", "— → DRAFT", "None", "None"),
        ("2", "Confirm sales order", "SALES_STAFF", "DRAFT → CONFIRMED", "PICK_ORDER → INVENTORY_MANAGER", "New order to pick: {order_number} to INVENTORY_MANAGER"),
        ("3", "Pick items from warehouse", "INVENTORY_MANAGER", "pick task PENDING → PICKED", "None", "None"),
        ("4", "Create package", "INVENTORY_MANAGER", "package DRAFT → PACKED", "None", "None"),
        ("5", "Commit fulfillment", "INVENTORY_MANAGER", "CONFIRMED → FULFILLED", "CREATE_INVOICE → SALES_STAFF", "SO fulfilled — create invoice to SALES_STAFF"),
        ("6", "Create and send invoice", "SALES_STAFF", "Invoice DRAFT → SENT", "None", "Invoice sent to TENANT_ADMIN"),
        ("7", "Cancel (alternative)", "SALES_STAFF / TENANT_ADMIN", "ANY → CANCELLED", "Cancel all open tasks", "Order cancelled to TENANT_ADMIN + SALES_STAFF"),
    ], [1.0, 4.6, 2.2, 3.3, 3.2, 3.3], font_size=8.4)
    add_paragraph(doc, "If a sales order is only partially fulfilled, the order moves to PARTIALLY_FULFILLED until all items are fulfilled.", italic=True)

    add_heading(doc, "7.2 Purchase Order Workflow", 2)
    build_table(doc, ["Step", "Action", "Actor", "Status Change", "Task Created", "Notification Sent"], [
        ("1", "Create purchase order", "PURCHASE_STAFF", "— → DRAFT", "None", "None"),
        ("2a", "Submit (low value)", "PURCHASE_STAFF", "DRAFT → SUBMITTED", "None", "None"),
        ("2b", "Submit (high value)", "PURCHASE_STAFF", "DRAFT → SUBMITTED", "APPROVE_PO → TENANT_ADMIN", "PO awaiting approval: {po_number} to TENANT_ADMIN"),
        ("3", "Approve (if high value)", "TENANT_ADMIN", "— (task completed)", "None", "None"),
        ("4", "Commit receipt", "PURCHASE_STAFF / INVENTORY_MANAGER", "SUBMITTED → PARTIALLY_RECEIVED", "PUTAWAY_STOCK → INVENTORY_MANAGER", "Stock received — putaway required to INVENTORY_MANAGER"),
        ("5", "Complete putaway", "INVENTORY_MANAGER", "putaway PENDING → COMPLETED", "RECORD_BILL → PURCHASE_STAFF", "Putaway complete — record bill to PURCHASE_STAFF"),
        ("6", "Record vendor bill", "PURCHASE_STAFF", "Bill DRAFT → RECORDED", "None", "None"),
    ], [1.0, 4.8, 2.3, 3.1, 3.2, 3.0], font_size=8.4)

    add_heading(doc, "7.3 Returns Workflow", 2)
    build_table(doc, ["Step", "Action", "Actor", "Status Change", "Task Created", "Notification Sent"], [
        ("1", "Create return", "SALES_STAFF", "— → DRAFT", "None", "None"),
        ("2", "Submit return", "SALES_STAFF", "DRAFT → SUBMITTED", "RETURN_QC → INVENTORY_MANAGER", "Return submitted — QC required to INVENTORY_MANAGER"),
        ("3", "Inspect items", "INVENTORY_MANAGER", "SUBMITTED → INSPECTION_PENDING", "None", "None"),
        ("4", "Process return", "INVENTORY_MANAGER", "INSPECTION_PENDING → PROCESSED", "None", "Return processed to SALES_STAFF + TENANT_ADMIN"),
    ], [1.0, 4.7, 2.4, 3.1, 3.1, 3.2], font_size=8.4)
    build_table(doc, ["QC Outcome", "Stock Effect", "Use When"], [
        ("ACCEPTED_RESTOCK", "Increases available stock", "Item is in good condition and can be resold"),
        ("ACCEPTED_BLOCKED", "Creates blocked stock record", "Item needs review before resale"),
        ("DAMAGED", "Creates blocked damaged stock", "Item is damaged and cannot be sold"),
        ("SCRAPPED", "Written off, no stock recovery", "Item is beyond use"),
        ("REJECTED", "No stock movement", "Return rejected, item is not accepted back"),
    ], [3.6, 4.8, 7.1], font_size=8.6)

    add_heading(doc, "7.4 Cycle Count Workflow", 2)
    build_table(doc, ["Step", "Action", "Actor", "Status Change", "Task Created", "Notification Sent"], [
        ("1", "Create session", "INVENTORY_MANAGER", "— → DRAFT", "None", "None"),
        ("2", "Add product lines", "INVENTORY_MANAGER", "DRAFT (in progress)", "None", "None"),
        ("3", "Count items physically", "INVENTORY_MANAGER", "DRAFT", "None", "None"),
        ("4", "Submit session", "INVENTORY_MANAGER", "DRAFT → SUBMITTED", "None", "None"),
        ("5", "Review variances", "INVENTORY_MANAGER", "SUBMITTED", "None", "None"),
        ("6", "Reconcile", "INVENTORY_MANAGER", "SUBMITTED → RECONCILED", "None", "None"),
    ], [1.0, 4.4, 2.5, 3.0, 3.0, 3.0], font_size=8.4)
    add_paragraph(doc, "Reconciliation re-snapshots the current system quantity at reconcile time, not at line-add time, so the final variance is accurate. A deterministic idempotency key per session and line prevents double-adjustment on retry.", italic=True)
    add_heading(doc, "7.5 Low Stock Reorder", 2)
    add_paragraph(doc, "When available stock drops below a product's reorder_level, the low_stock_check job automatically creates a REORDER_STOCK workflow task for PURCHASE_STAFF. A duplicate guard prevents creating a second task if one is already OPEN for the same product. The Reorder Suggestions report shows all products needing reorder ranked by urgency.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_page_break(doc)


def section_api(doc):
    add_heading(doc, "8. API Reference Summary", 1)
    add_paragraph(doc, "The complete API is available at GET /api/health. All endpoints require a Bearer token except /api/auth/login and /api/auth/register. The standard error envelope is { error: { code: string, message: string, request_id: string } }.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    api_groups = [
        ("Auth (/api/auth/)", [
            ("POST /login", "Public", "Email + password → access token + refresh token"),
            ("POST /register", "Public", "Create new tenant + admin user"),
            ("POST /refresh", "Bearer", "Rotate access token using refresh token"),
            ("GET /me", "Bearer", "Current user context"),
            ("POST /logout", "Bearer", "Revoke refresh token"),
            ("POST /forgot-password", "Public", "Request OTP reset code (always 204)"),
            ("POST /verify-reset-code", "Public", "Validate OTP → reset_token"),
            ("POST /reset-password", "Public", "Set new password with reset_token"),
        ]),
        ("Sales (/api/sales-orders/, /api/sales-fulfillments/)", [
            ("GET /sales-orders", "All sales roles", "List SO with filters"),
            ("POST /sales-orders", "Sales roles", "Create SO"),
            ("POST /sales-orders/{id}/confirm", "Sales roles", "Confirm SO, reserve stock, create PICK_ORDER task"),
            ("POST /sales-orders/{id}/cancel", "Sales + Admin", "Cancel SO, release stock, cancel tasks"),
            ("GET /sales-fulfillments", "Sales roles", "List all fulfillments"),
            ("POST /sales-orders/{id}/fulfillments", "Inventory", "Create fulfillment"),
            ("POST /fulfillments/{id}/commit", "Inventory", "Commit fulfillment, deduct stock"),
        ]),
        ("Purchase (/api/purchases/)", [
            ("GET /purchases", "Purchase roles", "List POs"),
            ("POST /purchases", "Purchase + Admin", "Create PO"),
            ("POST /purchases/{id}/submit", "Purchase roles", "Submit PO"),
            ("POST /purchases/{id}/receipts", "Purchase + Inventory", "Create receipt"),
            ("POST /purchase-receipts/{id}/commit", "Purchase + Inventory", "Commit receipt, add stock"),
        ]),
        ("Returns (/api/sales-returns/)", [
            ("POST /sales-returns", "Sales + Admin", "Create return"),
            ("POST /sales-returns/{id}/submit", "Sales + Admin", "Submit return, create RETURN_QC task"),
            ("POST /sales-returns/{id}/inspect", "Inventory + Admin", "QC inspection"),
            ("POST /sales-returns/{id}/process", "Inventory + Admin", "Process return, update stock"),
        ]),
        ("Inventory (/api/inventory/, /api/warehouses/, /api/cycle-counts/)", [
            ("GET /inventory/stock", "All roles", "Current stock levels"),
            ("POST /inventory/adjustments", "Inventory + Admin", "Manual stock adjustment"),
            ("GET /warehouses", "All roles", "List warehouses"),
            ("POST /cycle-counts", "Inventory", "Create count session"),
            ("POST /cycle-counts/{id}/submit", "Inventory", "Submit session"),
            ("POST /cycle-counts/{id}/reconcile", "Inventory", "Apply adjustments"),
        ]),
        ("Workflow (/api/workflow/)", [
            ("GET /workflow/my-tasks", "All except Viewer", "Tasks for current user's role"),
            ("GET /workflow/my-tasks/count", "All except Viewer", "Count of open tasks"),
            ("POST /workflow/tasks/{id}/start", "Role match", "OPEN → IN_PROGRESS"),
            ("POST /workflow/tasks/{id}/complete", "Role match", "→ COMPLETED"),
        ]),
        ("Documents (/api/invoices/, /api/bills/)", [
            ("POST /invoices", "Sales + Admin", "Create invoice (snapshots currency)"),
            ("POST /invoices/{id}/send", "Sales + Admin", "Send invoice PDF by email"),
            ("POST /bills", "Purchase + Admin", "Record vendor bill"),
        ]),
        ("AI Assistant (/api/faq/, /api/assistant/)", [
            ("GET /faq/suggestions", "All roles", "Suggested FAQ questions"),
            ("POST /faq/ask", "All roles", "Ask a question (RAG answer)"),
            ("POST /assistant/sessions", "Admin only", "Create copilot session"),
            ("POST /assistant/sessions/{id}/ask", "Admin only", "Ask the copilot (live data + RAG)"),
        ]),
    ]
    for title, rows in api_groups:
        add_heading(doc, title, 2)
        build_table(doc, ["Endpoint", "Access", "Description"], rows, [5.2, 3.2, 7.0], font_size=8.5)
    add_page_break(doc)


def section_ai(doc):
    add_heading(doc, "9. AI Copilot and FAQ", 1)
    add_heading(doc, "9.1 Overview", 2)
    add_paragraph(doc, "Warelyn has two AI features. The FAQ Assistant is available to all roles via a floating chat widget. The AI Copilot is available only to TENANT_ADMIN via a dedicated page. Both use Google Gemini for chat and embeddings, and MongoDB for vector storage.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "9.2 FAQ Assistant", 2)
    add_bullets(doc, [
        "Available to all roles via the floating widget in the bottom-right corner.",
        "Uses RAG, so answers are grounded only in the knowledge base.",
        "Hybrid search combines BM25-style lexical scoring with semantic cosine similarity.",
        "The knowledge base is made from source documents covering workflows, roles, reports, and troubleshooting.",
        "Answers below the configured confidence threshold are abstained.",
        "The assistant includes an application-only guardrail and refuses unrelated questions.",
        "Chat history persists in sessionStorage for the browser session.",
    ])

    add_heading(doc, "9.3 AI Copilot", 2)
    add_bullets(doc, [
        "Exclusive to TENANT_ADMIN.",
        "Extends the FAQ assistant with live tenant data queries.",
        "Uses natural language intent detection to turn questions into report queries and filters.",
        "Supports report views such as warehouse stock, low stock, reorder suggestions, stock movement, blocked stock, batch expiry, reconciliation, product valuation, open sales orders, pending receipts, open workflow tasks, and inventory summary.",
        "Can help draft workflows and operational summaries without mutating data.",
        "All queries are tenant-scoped.",
        "Report data renders as an inline table with insight bullets inside the chat bubble.",
    ])
    build_callout(doc, [
        "If the copilot always returns confidence 0.35, the GEMINI_API_KEY environment variable is not set. 0.35 is a hardcoded fallback, not a real retrieval score. Set WARELYN_GEMINI_API_KEY in your .env file to enable the full AI pipeline.",
    ], fill=LIGHT_AMBER)
    add_page_break(doc)


def section_dashboards(doc):
    add_heading(doc, "10. Dashboards", 1)
    add_heading(doc, "10.1 Role-Specific Dashboards", 2)
    add_paragraph(doc, "Every role gets a personalised dashboard with data relevant only to their function. The dashboard page detects the logged-in user's role and renders the appropriate sub-dashboard. All charts have independent filter controls, so changing one chart's filter does not affect the others.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    build_table(doc, ["Role", "Key KPIs", "Charts", "Unique Widgets"], [
        ("TENANT_ADMIN", "Revenue MTD, Spend MTD, Gross Margin, Open Tasks", "Revenue vs Spend dual-line, Team Task Donut, Order Health Bar, Stock Health", "Business Health Score, What needs attention alert list"),
        ("INVENTORY_MANAGER", "Low Stock, Blocked Stock, Expiring 30d, Stock Health Score", "Inbound vs Outbound Area, Warehouse Utilisation Bar, Low Stock by Category Donut, Movement Velocity Bar", "Expiry Timeline, Dead Stock Alert, Stock Accuracy Meter"),
        ("SALES_STAFF", "Revenue MTD, Orders Confirmed, Avg Order Value, Overdue Invoices", "Revenue by Day Line, Order Status Funnel, Top Customers Bar, Top Products Bar", "At-Risk Orders list, Invoice Aging table"),
        ("PURCHASE_STAFF", "Spend MTD, Pending Receipts, Overdue Bills, Avg Lead Time", "Spend by Day Line, Top Vendors Bar, PO Status Donut, Lead Time Distribution", "Reorder Urgency Score list, Pending Bills alert"),
        ("VIEWER", "Total Products, Stock Value, Low Stock Count, Open Orders", "Stock Movements Area, Order Status Bar", "Read-only — no actions"),
        ("SUPER_ADMIN", "Total Tenants, Active Tenants, New This Month, Audit Events 24h", "Tenant Growth 12-month, Audit Activity 30-day", "Most Active Tenants table, New Tenant Activation Tracker, Platform Health panel"),
    ], [2.1, 4.1, 6.0, 4.2], font_size=8.5)
    add_heading(doc, "10.2 Chart Filter System", 2)
    add_paragraph(doc, "Every chart manages its own independent filter state. Common filter options are date range (7d / 30d / 90d), warehouse dropdown, and a compare to previous period toggle. Date range filtering works by slicing the already-fetched data array, so changing chart filters does not make additional API calls.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_page_break(doc)


def section_currency(doc):
    add_heading(doc, "11. Multi-Currency Support", 1)
    for para in [
        "Warelyn supports 30+ ISO 4217 currencies. Each tenant sets a base currency in Settings → Company.",
        "Currency is snapshotted on invoice and bill creation, so the document's currency never changes even if the tenant later changes their setting. Historical documents always display their creation-time currency.",
        "formatMoney(value, currencyCode) always requires an explicit currency code. TenantSettingsContext provides the current currency to all components.",
    ]:
        add_paragraph(doc, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    build_table(doc, ["Code", "Currency", "Symbol"], [
        ("USD", "US Dollar", "$"),
        ("EUR", "Euro", "€"),
        ("GBP", "British Pound", "£"),
        ("INR", "Indian Rupee", "₹"),
        ("AED", "UAE Dirham", "د.إ"),
        ("SGD", "Singapore Dollar", "S$"),
        ("CNY", "Chinese Yuan", "¥"),
        ("BRL", "Brazilian Real", "R$"),
    ], [2.0, 8.0, 6.0], font_size=9.2)
    add_paragraph(doc, "The full supported list contains 30+ currencies and lives in the frontend currency registry.", italic=True)
    add_page_break(doc)


def section_notifications(doc):
    add_heading(doc, "12. Notifications", 1)
    add_paragraph(doc, "Warelyn sends in-app notifications when business events occur. Notifications are role-targeted — each notification goes only to users with the relevant role in the tenant. The bell icon in the top bar shows an unread count badge.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    build_table(doc, ["Trigger", "Target Role(s)", "Type", "Message Example"], [
        ("Sales order confirmed", "INVENTORY_MANAGER", "INFO", "New order to pick: SO-20250601-001"),
        ("Sales order cancelled", "TENANT_ADMIN, SALES_STAFF", "WARNING", "Order SO-20250601-001 cancelled"),
        ("Fulfillment committed", "SALES_STAFF", "INFO", "Order SO-20250601-001 fulfilled — create invoice"),
        ("Receipt committed", "INVENTORY_MANAGER", "INFO", "Stock received — putaway required for PO-001"),
        ("Putaway completed", "PURCHASE_STAFF", "INFO", "Putaway complete — record vendor bill"),
        ("Return submitted", "INVENTORY_MANAGER", "WARNING", "Return RTN-001 submitted — QC required"),
        ("Return processed", "SALES_STAFF, TENANT_ADMIN", "SUCCESS", "Return RTN-001 processed"),
        ("High-value PO submitted", "TENANT_ADMIN", "WARNING", "PO PO-001 awaiting your approval"),
        ("Low stock detected", "PURCHASE_STAFF", "WARNING", "Low stock: Product Name below reorder level"),
        ("User role changed", "Affected user", "INFO", "Your role has been changed to Inventory Manager"),
    ], [3.0, 4.0, 2.0, 7.0], font_size=8.3)
    add_paragraph(doc, "Notifications support mark-as-read (individual and all), clear (individual and all), and tab filtering (All / Unread / Cleared). Changes are persisted immediately with optimistic UI updates.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_page_break(doc)


def section_deployment(doc):
    add_heading(doc, "13. Deployment", 1)
    add_heading(doc, "13.1 Local Development Setup", 2)
    add_numbers(doc, [
        "Clone the repository.",
        "Copy backend/.env.example to backend/.env and fill in required values.",
        "Copy frontend/.env.example to frontend/.env.",
        "Run docker compose up --build -d.",
        "The backend runs Alembic migrations on startup, then seeds the super admin.",
        "Set WARELYN_SEED_ON_STARTUP=true to seed the demo tenant on first run.",
        "Frontend: cd frontend && npm install && npm run dev.",
        "Access the app at http://localhost:5173.",
        "API docs at http://localhost:8000/docs.",
    ])

    add_heading(doc, "13.2 Required Environment Variables", 2)
    build_table(doc, ["Variable", "Required", "Description"], [
        ("WARELYN_DATABASE_URL", "Yes", "MySQL connection string"),
        ("WARELYN_JWT_SECRET_KEY", "Yes", "Secure random string for JWT signing"),
        ("WARELYN_SUPER_ADMIN_EMAIL", "Yes", "Super admin login email"),
        ("WARELYN_SUPER_ADMIN_PASSWORD", "Yes", "Super admin login password"),
        ("WARELYN_GEMINI_API_KEY", "Yes (for AI)", "Google Gemini API key — without this, AI confidence is 0.35"),
        ("WARELYN_MONGO_URI", "Yes (for AI)", "MongoDB connection — stores RAG vector chunks"),
        ("WARELYN_SMTP_HOST", "No", "SMTP server host (defaults to MailHog)"),
        ("WARELYN_EMAIL_DELIVERY_MODE", "No", "mailhog / smtp / log"),
        ("WARELYN_SEED_ON_STARTUP", "No", "true/false — run demo seed on startup"),
    ], [4.0, 2.8, 9.2], font_size=8.4)
    build_callout(doc, [
        "Never use WARELYN_JWT_SECRET_KEY=secret in production. Generate a secure random string: python -c 'import secrets; print(secrets.token_hex(32))'",
    ], fill=LIGHT_RED)

    add_heading(doc, "13.3 Verification", 2)
    code = doc.add_paragraph()
    set_paragraph_format(code, before=0, after=0, line=1.0)
    run = code.add_run(
        "# Backend — must pass before any commit\n"
        "cd backend\n"
        ".venv/bin/python -m compileall app\n"
        ".venv/bin/python -m pytest -q\n\n"
        "# Frontend — must pass before any commit\n"
        "cd frontend\n"
        "npm run build\n\n"
        "# Full validation (Docker + DB available)\n"
        "./scripts/validate.sh"
    )
    set_run_font(run, name="Consolas", size=10.5, color=TEXT)
    add_page_break(doc)


def section_standards(doc):
    add_heading(doc, "14. Development Standards", 1)
    add_heading(doc, "14.1 Git Commit Convention", 2)
    add_paragraph(doc, "All commits follow Conventional Commits format: type(scope): imperative summary", bold=True)
    build_table(doc, ["Type", "Meaning"], [
        ("feat", "New feature or capability"),
        ("fix", "Bug fix"),
        ("refactor", "Code change, no behaviour change"),
        ("test", "Tests only"),
        ("docs", "Documentation only"),
        ("chore", "Dependencies, tooling, config"),
        ("security", "Security fix or hardening"),
        ("seed", "Seed data changes"),
        ("migration", "Database migration only"),
    ], [3.0, 13.0], font_size=9.0)
    add_paragraph(doc, "Examples from this project:", bold=True)
    for example in [
        "feat(workflow): add task auto-complete on fulfillment commit",
        "fix(returns): remove SALES_STAFF from QC inspect route",
        "fix(notifications): add db.commit() after all mutation methods",
        "security(rbac): add missing require_roles to cycle-counts cancel endpoint",
        "feat(copilot): add natural language query parameter extraction",
        "seed(dmart): add D-Mart wholesale tenant with multi-category demo data",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_format(p, before=0, after=1, line=1.08)
        set_run_font(p.add_run(example), name="Consolas", size=9.8)

    add_heading(doc, "14.2 The 10 Never-Violate Backend Rules", 2)
    rules = [
        "Stock mutations ONLY through InventoryEngine — never bypass it",
        "tenant_id from UserContext ONLY — never from request body",
        "Jinja2 for all template rendering — never str.format_map()",
        "Alembic migration for every DB schema change — never Base.metadata.create_all()",
        "db.commit() ONLY in service layer — never in repositories, routers, or models",
        "require_roles() on every endpoint — no unguarded endpoint ever",
        "Every repository query MUST filter by tenant_id — cross-tenant access is a critical security bug",
        "AppError is the only exception class — never raise raw HTTPException in services",
        "Business logic NEVER in routers — routers translate HTTP only",
        "Workflow side effects (tasks, notifications) in try/except — they must never break the primary operation",
    ]
    for idx, rule in enumerate(rules, start=1):
        p = doc.add_paragraph()
        set_paragraph_format(p, before=0, after=2, line=1.08)
        run = p.add_run(f"{idx}. ")
        set_run_font(run, bold=True, color="E53935")
        run2 = p.add_run(rule)
        set_run_font(run2)

    add_heading(doc, "14.3 Code Quality Checklist", 2)
    checklist_items = [
        "Type hints on all function signatures",
        "Decimal for all monetary values (never float)",
        "datetime always timezone-aware (datetime.now(timezone.utc))",
        "Unique constraints on natural keys per tenant",
        "Indexes on (tenant_id, status) for hot query paths",
        "Every new endpoint has 3 tests: happy path, 401, 403",
        "Every mutation tested for persistence (fetch after write)",
        "Tenant isolation tested (cross-tenant 403/404)",
        "Frontend: controlled components only (no getElementById)",
        "Frontend: every list has loading state + empty state",
        "Frontend: optimistic UI updates with revert on failure",
    ]
    build_table(doc, ["☐", "Checklist Item"], [("☐", item) for item in checklist_items], [0.8, 15.2], font_size=9.0)
    add_page_break(doc)


def section_quick_reference(doc):
    add_heading(doc, "15. Quick Reference Cards", 1)
    add_heading(doc, "15.1 Workflow Task Reference Card", 2)
    build_table(doc, ["Step Key", "Created By", "Assigned To", "Triggered When", "Auto-Completes When"], [
        ("PICK_ORDER", "confirm_sales_order", "INVENTORY_MANAGER", "SO confirmed", "All items picked"),
        ("PUTAWAY_STOCK", "commit_purchase_receipt", "INVENTORY_MANAGER", "Receipt committed", "Putaway marked complete"),
        ("RECORD_BILL", "complete_putaway_task", "PURCHASE_STAFF", "Putaway complete", "Bill recorded"),
        ("CREATE_INVOICE", "commit_fulfillment", "SALES_STAFF", "Fulfillment committed", "Invoice sent"),
        ("RETURN_QC", "submit_return", "INVENTORY_MANAGER", "Return submitted", "Return processed"),
        ("APPROVE_PO", "submit_purchase_order", "TENANT_ADMIN", "High-value PO submitted", "PO approved"),
        ("REORDER_STOCK", "low_stock_check job", "PURCHASE_STAFF", "Stock below reorder level", "PO created for that product"),
    ], [2.2, 3.0, 3.0, 3.2, 4.0], font_size=8.5)

    add_heading(doc, "15.2 Status Reference Card", 2)
    status_sections = {
        "Sales Order Status": [
            ("DRAFT", "Created, not confirmed, stock not reserved"),
            ("CONFIRMED", "Stock reserved, pick task created"),
            ("PARTIALLY_FULFILLED", "Some items fulfilled, others pending"),
            ("FULFILLED", "All items fulfilled, invoice task created"),
            ("CLOSED", "Invoice sent and paid"),
            ("CANCELLED", "Cancelled, stock released, tasks cancelled"),
        ],
        "Purchase Order Status": [
            ("DRAFT", "Created, not sent to vendor"),
            ("SUBMITTED", "Sent to vendor, awaiting delivery"),
            ("PARTIALLY_RECEIVED", "Some items received"),
            ("RECEIVED", "All items received"),
            ("CANCELLED", "Cancelled"),
        ],
        "Return Status": [
            ("DRAFT", "Created, not submitted"),
            ("SUBMITTED", "Submitted, QC task created"),
            ("INSPECTION_PENDING", "QC in progress"),
            ("PROCESSED", "QC complete, stock updated"),
            ("CANCELLED", "Cancelled"),
        ],
        "Workflow Task Status": [
            ("OPEN", "Created, waiting to be started"),
            ("IN_PROGRESS", "User has claimed and started the task"),
            ("COMPLETED", "Work done, auto-completed or manually completed"),
            ("CANCELLED", "Parent entity cancelled, task no longer relevant"),
        ],
    }
    for title, rows in status_sections.items():
        add_heading(doc, title, 3)
        build_table(doc, ["Status", "Meaning"], rows, [3.2, 12.8], font_size=8.8)

    add_heading(doc, "15.3 API Error Codes Reference", 2)
    build_table(doc, ["Error Code", "HTTP", "Meaning"], [
        ("MISSING_TOKEN", "401", "No Bearer token in request"),
        ("INVALID_TOKEN", "401", "Token expired or tampered"),
        ("FORBIDDEN_ROLE", "403", "User's role not in allowed roles"),
        ("TENANT_ACCESS_DENIED", "403", "Tenant context required"),
        ("INSUFFICIENT_STOCK", "409", "Not enough available stock to reserve or deduct"),
        ("INVALID_RESET_CODE", "400", "OTP code wrong or expired"),
        ("RESET_TOKEN_ALREADY_USED", "400", "Reset link already consumed"),
        ("DUPLICATE_SKU", "409", "SKU already exists for this tenant"),
        ("PACK_ORDER_ALREADY_COMPLETED", "409", "Tried to pack items already in a package"),
        ("RETURN_QUANTITY_EXCEEDS_FULFILLED", "400", "Returning more than was fulfilled"),
        ("CYCLE_COUNT_SESSION_NOT_FOUND", "404", "Session does not exist for this tenant"),
    ], [4.0, 1.6, 10.4], font_size=8.5)


def main():
    doc = Document()
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings.element.append(update_fields)
    configure_sections(doc)
    apply_document_styles(doc)
    add_header_footer(doc.sections[0])
    add_cover_page(doc)
    add_toc_page(doc)
    section_overview(doc)
    section_architecture(doc)
    section_database(doc)
    section_roles(doc)
    section_workflows(doc)
    section_api(doc)
    section_ai(doc)
    section_dashboards(doc)
    section_currency(doc)
    section_notifications(doc)
    section_deployment(doc)
    section_standards(doc)
    section_quick_reference(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Generated: {OUT_PATH.relative_to(REPO_ROOT)}")


if __name__ == "__main__":
    main()
